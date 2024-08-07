import fnmatch
import glob
import os
import re
import subprocess
import zipfile
from pathlib import Path

import cv2
import docx
import img2pdf
import numpy as np
import pandas as pd
import pdfplumber
import pytesseract
import streamlit as st
from docx.api import Document

# from paddleocr import PaddleOCR
from pdf2image import convert_from_path
from PIL import Image

Image.MAX_IMAGE_PIXELS = None

# uploadpath = "uploads/"

# ocr = PaddleOCR(use_angle_cls=True, lang="ch")


def docxurl2txt(url):
    text = ""
    try:
        doc = docx.Document(url)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
            text = "\n".join(fullText)
    except Exception as e:
        st.error(str(e))

    return text


def pdfurl2txt(url):
    #     response = requests.get(url)
    #     source_stream = BytesIO(response.content)
    result = ""
    try:
        #         with pdfplumber.open(source_stream) as pdf:
        with pdfplumber.open(url) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt != "":
                    result += txt
    except Exception as e:
        st.error(str(e))
    return result


# def paddleocr2text(image_file):
#     result = ocr.ocr(image_file, cls=True)
#     # print(result)
#     txtls = [line[1][0] for line in result]
#     #     print(txtls)
#     txt = "\n".join(txtls)
#     return txt


def pytesseract2text(image_file):
    text = pytesseract.image_to_string(image_file, lang="chi_sim")
    return text


def pdfurl2ocr(url, uploadpath):
    PDF_file = Path(url)
    # Store all the pages of the PDF in a variable
    image_file_list = []
    text = ""
    # with TemporaryDirectory() as tempdir:
    pdf_pages = convert_from_path(PDF_file, 500)
    # Iterate through all the pages stored above
    for page_enumeration, page in enumerate(pdf_pages, start=1):
        # enumerate() "counts" the pages for us.

        # Create a file name to store the image
        filename = os.path.join(uploadpath, "page_" + str(page_enumeration) + ".jpg")

        # Save the image of the page in system
        page.save(filename, "JPEG")
        image_file_list.append(filename)

    # Iterate from 1 to total number of pages
    for image_file in image_file_list:
        # text += paddleocr2text(image_file)
        text += pytesseract2text(image_file)
        # delete image file
        os.remove(image_file)

    return text


def docxurl2ocr(url, uploadpath):
    z = zipfile.ZipFile(url)
    all_files = z.namelist()
    images = sorted(filter(lambda x: x.startswith("word/media/"), all_files))

    # Store all the pages of the PDF in a variable
    image_file_list = []
    text = ""
    # with TemporaryDirectory() as tempdir:
    # Iterate through all the pages stored above
    for page_enumeration, image in enumerate(images):
        # enumerate() "counts" the pages for us.
        img = z.open(image).read()
        # Create a file name to store the image
        filename = os.path.basename(image)
        filepath = os.path.join(uploadpath, filename)
        #             print(filename)
        # Save the image of the page in system
        f = open(filepath, "wb")
        f.write(img)
        image_file_list.append(filepath)

    # Iterate from 1 to total number of pages
    for image_file in image_file_list:
        # text += paddleocr2text(image_file)
        text += pytesseract2text(image_file)
        # delete image file
        os.remove(image_file)

    return text


def picurl2ocr(url):
    text = ""
    # text += paddleocr2text(url)
    text += pytesseract2text(url)
    return text


def find_files(path: str, glob_pat: str, ignore_case: bool = False):
    rule = (
        re.compile(fnmatch.translate(glob_pat), re.IGNORECASE)
        if ignore_case
        else re.compile(fnmatch.translate(glob_pat))
    )
    return [
        n for n in glob.glob(os.path.join(path, "*.*"), recursive=True) if rule.match(n)
    ]


def save_uploadedfile(uploadedfile, uploadpath):
    with open(os.path.join(uploadpath, uploadedfile.name), "wb") as f:
        f.write(uploadedfile.getbuffer())
    return st.success("上传文件:{} 成功。".format(uploadedfile.name))


def docxconvertion(uploadpath):
    docdest = os.path.join(uploadpath, "doc")
    wpsdest = os.path.join(uploadpath, "wps")
    docxdest = os.path.join(uploadpath, "docx")
    docmdest = os.path.join(uploadpath, "docm")

    # create folder if not exist
    if not os.path.exists(docdest):
        os.makedirs(docdest)
    if not os.path.exists(wpsdest):
        os.makedirs(wpsdest)
    if not os.path.exists(docxdest):
        os.makedirs(docxdest)
    if not os.path.exists(docmdest):
        os.makedirs(docmdest)

    docfiles = find_files(uploadpath, "*.doc", True)
    wpsfiles = find_files(uploadpath, "*.wps", True)
    docxfiles = find_files(uploadpath, "*.docx", True)
    docmfiles = find_files(uploadpath, "*.docm", True)

    for filepath in docfiles:
        st.info(filepath)
        # filename = os.path.basename(filepath)
        #     print(filename)
        #         output = subprocess.check_output(["soffice","--headless","--convert-to","docx",file,"--outdir",dest])
        subprocess.call(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                filepath,
                "--outdir",
                docdest,
            ]
        )

    for filepath in wpsfiles:
        st.info(filepath)
        # filename = os.path.basename(filepath)
        #     print(filename)
        #         output = subprocess.check_output(["soffice","--headless","--convert-to","docx",file,"--outdir",dest])
        subprocess.call(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                filepath,
                "--outdir",
                wpsdest,
            ]
        )

    # for filepath in doccfiles:
    #     print (filepath)
    #     filename=os.path.basename(filepath)
    # #     print(filename)
    # #         output = subprocess.check_output(["soffice","--headless","--convert-to","docx",file,"--outdir",dest])
    #     subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filepath,"--outdir",doccdest])

    for filepath in docxfiles:
        st.info(filepath)
        # filename = os.path.basename(filepath)
        #     print(filename)
        #         output = subprocess.check_output(["soffice","--headless","--convert-to","docx",file,"--outdir",dest])
        subprocess.call(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                filepath,
                "--outdir",
                docxdest,
            ]
        )

    for filepath in docmfiles:
        st.info(filepath)
        # filename = os.path.basename(filepath)
        #     print(filename)
        #         output = subprocess.check_output(["soffice","--headless","--convert-to","docx",file,"--outdir",dest])
        subprocess.call(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                filepath,
                "--outdir",
                docmdest,
            ]
        )


def get_uploadfiles(uploadpath):
    fileslist = glob.glob(uploadpath + "/*.*", recursive=True)
    basenamels = []
    for file in fileslist:
        basenamels.append(os.path.basename(file))
    return basenamels


def remove_uploadfiles(uploadpath):
    files = glob.glob(uploadpath + "**/*.*", recursive=True)

    for f in files:
        try:
            os.remove(f)
        except OSError as e:
            st.error("Error: %s : %s" % (f, e.strerror))


# convert all files in uploadfolder to text
def convert_uploadfiles(txtls, uploadpath):
    resls = []
    for file in txtls:
        # st.info(file)
        try:
            # datapath=file
            datapath = os.path.join(uploadpath, file)
            #     get file ext
            base, ext = os.path.splitext(file)

            if ext.lower() == ".doc":
                # datapath = uploadpath + "doc/" + base + ".docx"
                datapath = os.path.join(uploadpath, "doc", base + ".docx")
                st.info(datapath)
                text = docxurl2txt(datapath)
                text1 = text.translate(str.maketrans("", "", r" \n\t\r\s"))
                if text1 == "":
                    text = docxurl2ocr(datapath, uploadpath)

            elif ext.lower() == ".wps":
                # datapath = uploadpath + "wps/" + base + ".docx"
                datapath = os.path.join(uploadpath, "wps", base + ".docx")
                st.info(datapath)
                text = docxurl2txt(datapath)
                text1 = text.translate(str.maketrans("", "", r" \n\t\r\s"))
                if text1 == "":
                    text = docxurl2ocr(datapath, uploadpath)

            #         elif ext.lower()=='doc.docx':
            #             datapath=os.path.join(filepath,'docc',file)
            #             print(datapath)
            #             text=docxurl2txt(datapath)
            elif ext.lower() == ".docx":
                st.info(datapath)
                text = docxurl2txt(datapath)
                text1 = text.translate(str.maketrans("", "", r" \n\t\r\s"))
                if text1 == "":
                    datapath = os.path.join(uploadpath, "docx", file)
                    st.info(datapath)
                    text = docxurl2txt(datapath)
                    text2 = text.translate(str.maketrans("", "", r" \n\t\r\s"))
                    if text2 == "":
                        text = docxurl2ocr(datapath, uploadpath)

            elif ext.lower() == ".pdf":
                text = pdfurl2txt(datapath)
                text1 = text.translate(str.maketrans("", "", r" \n\t\r\s"))
                if text1 == "":
                    text = pdfurl2ocr(datapath, uploadpath)

            elif (
                ext.lower() == ".png"
                or ext.lower() == ".jpg"
                or ext.lower() == ".jpeg"
                or ext.lower() == ".bmp"
                or ext.lower() == ".tiff"
            ):
                text = picurl2ocr(datapath)
            else:
                text = np.nan
        except Exception as e:
            st.error(str(e))
            text = ""
        resls.append(text)
    return resls


# extract text from files
def extract_text(df, uploadpath):
    txtls = df["文件"].tolist()
    resls = convert_uploadfiles(txtls, uploadpath)
    df["文本"] = resls
    return df


def picurl2table(url):
    image, mylistx, mylisty = seg_pic(url)
    # display(image)
    # st.image(image)
    mylist = table_ocr(image, mylistx, mylisty)
    df = pd.DataFrame(mylist)
    return df


def table_ocr(image, mylistx, mylisty):
    # ocr = PaddleOCR(det=True)
    # 循环y坐标，x坐标分割表格
    mylist = []
    for i in range(len(mylisty) - 1):
        row = []
        for j in range(len(mylistx) - 1):
            # 在分割时，第一个参数为y坐标，第二个参数为x坐标
            ROI = image[
                mylisty[i] + 3 : mylisty[i + 1] - 3, mylistx[j] : mylistx[j + 1] - 3
            ]  # 减去3的原因是由于我缩小ROI范围
            # cv2.imshow("分割后子图片展示：", ROI)
            # cv2.waitKey(0)
            # result = ocr.ocr(ROI, det=True)
            # text = paddleocr2text(ROI)
            text = pytesseract2text(ROI)
            # text_len = len(result)
            # tmptxt = " "
            # txt = " "
            # if text_len != 0:
            #     text = ""
            #     for idx in range(len(result)):
            #         res = result[idx]
            #         for line in res:
            #             tmptxt, _ = line[-1]
            #             txt = txt + "\n" + tmptxt
            #         text += txt
            row.append(text)
            j = j + 1
        i = i + 1
        mylist.append(row)

    return mylist


def seg_pic(img):
    image = cv2.imread(img, 1)

    # 灰度图片
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    # 二值化
    binary = cv2.adaptiveThreshold(
        ~gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 35, -5
    )
    # Calculate binary threshold using Otsu's method
    # _, binary = cv2.threshold(~gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # ret,binary = cv2.threshold(~gray, 127, 255, cv2.THRESH_BINARY)
    # cv2.imshow("二值化图片：", binary)  # 展示图片
    # cv2.waitKey(0)

    rows, cols = binary.shape

    # 识别横线
    # kernel_width = cols // 20
    # kernel_height = rows // 20
    # 用于扩充或者腐蚀图像
    # kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_width, kernel_height))

    scale = 40
    # 识别横线
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (cols // scale, 1))
    eroded = cv2.erode(binary, kernel, iterations=1)
    # cv2.imshow("Eroded Image",eroded)
    dilatedcol = cv2.dilate(eroded, kernel, iterations=1)
    # cv2.imshow("表格横线展示：", dilatedcol)
    # cv2.waitKey(0)

    # 识别竖线
    scale = 20
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, rows // scale))
    eroded = cv2.erode(binary, kernel, iterations=1)
    dilatedrow = cv2.dilate(eroded, kernel, iterations=1)
    # cv2.imshow("表格竖线展示：", dilatedrow)
    # cv2.waitKey(0)

    # 标识交点
    bitwiseAnd = cv2.bitwise_and(dilatedcol, dilatedrow)
    # cv2.imshow("表格交点展示：", bitwiseAnd)
    # cv2.waitKey(0)
    # cv2.imwrite("my.png",bitwiseAnd) #将二值像素点生成图片保存

    # 标识表格
    # merge = cv2.add(dilatedcol, dilatedrow)
    # cv2.imshow("表格整体展示：", merge)
    # cv2.waitKey(0)

    # 两张图片进行减法运算，去掉表格框线
    # merge2 = cv2.subtract(binary, merge)
    # cv2.imshow("图片去掉表格框线展示：", merge2)
    # cv2.waitKey(0)

    # 识别黑白图中的白色交叉点，将横纵坐标取出
    ys, xs = np.where(bitwiseAnd > 0)

    mylisty = []  # 纵坐标
    mylistx = []  # 横坐标

    # 通过排序，获取跳变的x和y的值，说明是交点，否则交点会有好多像素值值相近，我只取相近值的最后一点
    # 这个10的跳变不是固定的，根据不同的图片会有微调，基本上为单元格表格的高度（y坐标跳变）和长度（x坐标跳变）
    i = 0
    myxs = np.sort(xs)
    for i in range(len(myxs) - 1):
        if myxs[i + 1] - myxs[i] > 10:
            mylistx.append(myxs[i])
        i = i + 1
    if len(myxs) > 0:
        mylistx.append(myxs[i])  # 要将最后一个点加入

    i = 0
    myys = np.sort(ys)
    # print(np.sort(ys))
    for i in range(len(myys) - 1):
        if myys[i + 1] - myys[i] > 10:
            mylisty.append(myys[i])
        i = i + 1
    if len(myys) > 0:
        mylisty.append(myys[i])  # 要将最后一个点加入
    return image, mylistx, mylisty


def pdfurl2tableocr(url, uploadpath):
    PDF_file = Path(url)
    # Store all the pages of the PDF in a variable
    image_file_list = []
    # with TemporaryDirectory() as tempdir:
    pdf_pages = convert_from_path(PDF_file, 500)
    # Iterate through all the pages stored above
    for page_enumeration, page in enumerate(pdf_pages, start=1):
        # enumerate() "counts" the pages for us.

        # Create a file name to store the image
        filename = os.path.join(uploadpath, "page_" + str(page_enumeration) + ".jpg")

        # Save the image of the page in system
        page.save(filename, "JPEG")
        image_file_list.append(filename)

    resls = []
    # Iterate from 1 to total number of pages
    for image_file in image_file_list:
        st.image(image_file)
        df = picurl2table(image_file)
        # df = improved_table_ocr(image_file)
        resls.append(df)
        # delete image file
        os.remove(image_file)
    if resls:
        resdf = pd.concat(resls)
    else:
        resdf = pd.DataFrame()
    return resdf


def word2df(word):
    document = Document(word)
    data = []
    for table in document.tables:
        tb = []
        for row in table.rows:
            rl = []
            for cell in row.cells:
                #             print(cell.text)
                rl.append(cell.text)
            tb.append(rl)
        data.append(tb)
    dataarray = np.array(data)
    shape = dataarray.shape
    # st.write(shape)
    if len(shape) == 3:
        dataarray1 = dataarray.reshape(shape[0] * shape[1], shape[2])
        df = pd.DataFrame(dataarray1)
    elif len(shape) == 2:
        df = pd.DataFrame(dataarray)
    else:
        df = pd.DataFrame()
    return df
    # for table in document.tables:
    #     df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    #     for i, row in enumerate(table.rows):
    #         for j, cell in enumerate(row.cells):
    #             if cell.text:
    #                 df[i][j] = cell.text
    #     data.append(pd.DataFrame(df))
    # if data:
    #     resdf = pd.concat(data)
    # else:
    #     resdf = pd.DataFrame()
    # return resdf


def get_convertfname(file_name, temppath, ext):
    # get file extension
    extension = file_name.split(".")[-1]
    # get base name
    basename = file_name.split(".")[0]
    # converted word filename
    fname = basename + "." + ext
    st.write(fname)
    # get excel file content
    filepath = temppath + "/" + extension.lower() + "/" + fname
    return filepath


def docx2pdf(filepath, temppath):
    st.info(filepath)
    # get file extension
    extension = filepath.split(".")[-1]
    docdest = os.path.join(temppath, extension.lower())
    subprocess.call(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            filepath,
            "--outdir",
            docdest,
        ]
    )


def img_to_pdf(picpath, temppath):
    # get pdf file name
    pdfname = picpath.split("/")[-1].split(".")[0] + ".pdf"
    # get pdf file path
    pdfpath = os.path.join(temppath, pdfname)

    with open(pdfpath, "wb") as f:
        f.write(img2pdf.convert(picpath))


def preprocess_image(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (3, 3), 0)
    thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
    return thresh


def find_table_contours(thresh):
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    return sorted(contours, key=cv2.contourArea, reverse=True)


def get_table_rows(contours, image):
    rows = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        if (
            w > image.shape[1] * 0.5 and 20 < h < image.shape[0] * 0.1
        ):  # Filter for likely table rows
            rows.append((y, y + h))
    return sorted(rows)


def perform_ocr(image):
    return pytesseract.image_to_string(image, lang="chi_sim+eng", config="--psm 6")


def improved_table_ocr(image_path):
    image = cv2.imread(image_path)

    thresh = preprocess_image(image)
    contours = find_table_contours(thresh)
    rows = get_table_rows(contours, image)

    table_data = []
    header = []

    for i, (y1, y2) in enumerate(rows):
        row_image = image[y1:y2, :]
        row_text = perform_ocr(row_image).replace("\n", " ").strip()

        if i == 0:  # Assume first row is header
            header = row_text.split()
        elif row_text and len(row_text) > 10:  # Ignore very short texts
            table_data.append(row_text.split())

    # Ensure all rows have the same number of columns as the header
    max_cols = len(header)
    table_data = [row + [""] * (max_cols - len(row)) for row in table_data]

    df = pd.DataFrame(table_data, columns=header)
    return df
