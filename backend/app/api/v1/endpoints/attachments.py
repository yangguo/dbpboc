from typing import List, Optional, Dict, Any
import logging
import asyncio
import time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import unquote, urlparse
from fastapi import APIRouter, HTTPException, status, Query, BackgroundTasks, Depends
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
import os
import glob
import pandas as pd
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import tempfile
import json
import base64
import subprocess
import shutil
from pathlib import Path
from app.core.config import settings
import uuid

router = APIRouter()
logger = logging.getLogger("uvicorn.error")

# Constants
TEMP_PATH = "../temp"
PBOC_DATA_PATH = "../pboc"

# City mappings
org2name = {
    "天津": "tianjin", "重庆": "chongqing", "上海": "shanghai", "兰州": "lanzhou",
    "拉萨": "lasa", "西宁": "xining", "乌鲁木齐": "wulumuqi", "南宁": "nanning",
    "贵阳": "guiyang", "福州": "fuzhou", "成都": "chengdu", "呼和浩特": "huhehaote",
    "郑州": "zhengzhou", "北京": "beijing", "合肥": "hefei", "厦门": "xiamen",
    "海口": "haikou", "大连": "dalian", "广州": "guangzhou", "太原": "taiyuan",
    "石家庄": "shijiazhuang", "总部": "zongbu", "昆明": "kunming", "青岛": "qingdao",
    "沈阳": "shenyang", "长沙": "changsha", "深圳": "shenzhen", "武汉": "wuhan",
    "银川": "yinchuan", "西安": "xian", "哈尔滨": "haerbin", "长春": "changchun",
    "宁波": "ningbo", "杭州": "hangzhou", "南京": "nanjing", "济南": "jinan",
    "南昌": "nanchang",
}

class AttachmentItem(BaseModel):
    id: str
    link: str
    downloadUrl: str
    fileName: str
    status: str = 'pending'
    fileSize: Optional[str] = None

class FileItem(BaseModel):
    id: str
    fileName: str
    fileType: str
    link: str
    filePath: str
    status: str = 'pending'

class ProcessedDataItem(BaseModel):
    id: str
    data: dict

def create_session_with_retry() -> requests.Session:
    """Create a requests session with retry strategy and sane defaults.

    Note: Some PBOC regional sites do not respond to HEAD and can be slow
    to deliver binary payloads. We prioritize resilient GET behavior here.
    """
    session = requests.Session()
    retry_strategy = Retry(
        total=3,
        connect=3,
        read=3,
        backoff_factor=1,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "HEAD", "OPTIONS"],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry_strategy, pool_maxsize=10)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    # Avoid inheriting proxies which can cause connection resets to gov CN sites
    session.trust_env = False
    return session

def format_bytes(bytes_count: int) -> str:
    """Format bytes to human readable format"""
    if bytes_count < 1024:
        return f"{bytes_count} B"
    elif bytes_count < 1024 * 1024:
        return f"{bytes_count / 1024:.1f} KB"
    else:
        return f"{bytes_count / (1024 * 1024):.1f} MB"

def format_speed(bytes_per_second: float) -> str:
    """Format download speed"""
    if bytes_per_second < 1024:
        return f"{bytes_per_second:.1f} B/s"
    elif bytes_per_second < 1024 * 1024:
        return f"{bytes_per_second / 1024:.1f} KB/s"
    else:
        return f"{bytes_per_second / (1024 * 1024):.1f} MB/s"

def get_csvdf(folder: str, beginwith: str) -> pd.DataFrame:
    """Get CSV files matching pattern from folder"""
    files = glob.glob(os.path.join(folder, f"**/{beginwith}*.csv"), recursive=True)
    dflist = []
    for filepath in files:
        try:
            df = pd.read_csv(filepath)
            dflist.append(df)
        except Exception as e:
            logger.error(f"Error reading {filepath}: {e}")
    
    if dflist:
        result_df = pd.concat(dflist)
        result_df.reset_index(drop=True, inplace=True)
        return result_df
    return pd.DataFrame()

@router.get("/download-list/{org_name}")
async def get_download_list(org_name: str) -> List[AttachmentItem]:
    """Get list of attachments to download for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Look for download files in temp folder - check both locations
        beginwith = f"pboctodownload{org_name_index}"
        
        # First try the organization-specific subfolder
        org_temp_path = os.path.join(TEMP_PATH, org_name_index)
        df = get_csvdf(org_temp_path, beginwith)
        
        # If not found, try the main temp folder
        if df.empty:
            df = get_csvdf(TEMP_PATH, beginwith)
        
        if df.empty:
            return []
        
        # Check downloads directory for existing files
        downloads_dir = os.path.join(TEMP_PATH, org_name_index, "downloads")
        existing_files = set()
        if os.path.exists(downloads_dir):
            existing_files = set(os.listdir(downloads_dir))
        
        attachments = []
        for idx, row in df.iterrows():
            download_url = row.get('download', '')
            file_name = os.path.basename(download_url) if download_url else f"file_{idx}"
            # Decode URL-encoded filename
            from urllib.parse import unquote
            file_name = unquote(file_name)
            
            # Clean filename for comparison
            import re
            clean_filename = re.sub(r'[<>:"/\\|?*]', '_', file_name)
            
            # Check if file already exists
            file_exists = clean_filename in existing_files
            # Also check for files with counter suffix (filename_1.ext, filename_2.ext, etc.)
            if not file_exists:
                name, ext = os.path.splitext(clean_filename)
                for existing_file in existing_files:
                    if existing_file.startswith(name) and existing_file.endswith(ext):
                        # Check if it matches pattern filename_N.ext
                        middle_part = existing_file[len(name):-len(ext)] if ext else existing_file[len(name):]
                        if middle_part.startswith('_') and middle_part[1:].isdigit():
                            file_exists = True
                            break
            
            # Get file size if it exists
            file_size = None
            if file_exists:
                try:
                    file_path = os.path.join(downloads_dir, clean_filename)
                    if os.path.exists(file_path):
                        size_bytes = os.path.getsize(file_path)
                        if size_bytes < 1024:
                            file_size = f"{size_bytes} B"
                        elif size_bytes < 1024 * 1024:
                            file_size = f"{size_bytes / 1024:.1f} KB"
                        else:
                            file_size = f"{size_bytes / (1024 * 1024):.1f} MB"
                except:
                    pass
            
            attachment = AttachmentItem(
                id=str(idx),
                link=row.get('link', ''),
                downloadUrl=download_url,
                fileName=file_name,
                status='completed' if file_exists else 'pending',
                fileSize=file_size
            )
            attachments.append(attachment)
        
        return attachments
    
    except Exception as e:
        logger.error(f"Error getting download list for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get download list")

@router.get("/file-list/{org_name}")
async def get_file_list(org_name: str) -> List[FileItem]:
    """Get list of files for processing for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Look for file list in temp folder - check both locations
        beginwith = f"pboctofile{org_name_index}"
        
        # First try the organization-specific subfolder
        org_temp_path = os.path.join(TEMP_PATH, org_name_index)
        df = get_csvdf(org_temp_path, beginwith)
        
        # If not found, try the main temp folder
        if df.empty:
            df = get_csvdf(TEMP_PATH, beginwith)
        
        if df.empty:
            return []
        
        files = []
        for idx, row in df.iterrows():
            file_name = row.get('file', '')
            file_ext = os.path.splitext(file_name)[1].lower()
            
            # Decode URL-encoded filename
            from urllib.parse import unquote
            file_name = unquote(file_name)
            
            # Determine file type
            if file_ext in ['.xlsx', '.xls', '.et', '.ett']:
                file_type = 'excel'
            elif file_ext == '.pdf':
                file_type = 'pdf'
            elif file_ext in ['.docx', '.doc', '.wps', '.docm']:
                file_type = 'word'
            elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tif']:
                file_type = 'image'
            else:
                file_type = 'other'
            
            file_item = FileItem(
                id=str(idx),
                fileName=file_name,
                fileType=file_type,
                link=row.get('link', ''),
                filePath=os.path.join(TEMP_PATH, org_name_index, "downloads", file_name),
                status='pending'
            )
            files.append(file_item)
        
        return files
    
    except Exception as e:
        logger.error(f"Error getting file list for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get file list")

@router.get("/processed-data/{org_name}")
async def get_processed_data(org_name: str) -> List[ProcessedDataItem]:
    """Get processed table data for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Look for processed table data in temp folder - check both locations
        beginwith = f"pboctotable{org_name_index}"
        
        # First try the organization-specific subfolder
        org_temp_path = os.path.join(TEMP_PATH, org_name_index)
        df = get_csvdf(org_temp_path, beginwith)
        
        # If not found, try the main temp folder
        if df.empty:
            df = get_csvdf(TEMP_PATH, beginwith)
        
        if df.empty:
            return []
        
        # Get existing links from all pboccat files to exclude them
        existing_links = set()
        try:
            pboccat_df = get_csvdf(PBOC_DATA_PATH, "pboccat")
            if not pboccat_df.empty and 'id' in pboccat_df.columns:
                existing_links = set(pboccat_df['id'].dropna().tolist())
                logger.info(f"Found {len(existing_links)} existing links in pboccat files")
        except Exception as e:
            logger.warning(f"Error reading pboccat files: {e}")
        
        # Filter out rows with links that already exist in pboccat files
        if existing_links and 'link' in df.columns:
            original_count = len(df)
            df = df[~df['link'].isin(existing_links)]
            filtered_count = len(df)
            logger.info(f"Filtered processed data: {original_count} -> {filtered_count} (excluded {original_count - filtered_count} existing links)")
        
        processed_data = []
        for idx, (original_idx, row) in enumerate(df.iterrows(), 1):
            row_dict = row.to_dict()
            # 添加序号字段
            row_dict['序号'] = idx
            # 使用原始索引+link的hash作为稳定的ID，确保选择状态的一致性
            link = row_dict.get('link', '')
            stable_id = f"{original_idx}_{hash(link) % 10000}" if link else str(original_idx)
            data_item = ProcessedDataItem(
                id=stable_id,
                data=row_dict
            )
            processed_data.append(data_item)
        
        return processed_data
    
    except Exception as e:
        logger.error(f"Error getting processed data for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get processed data")

@router.post("/save-processed/{org_name}")
async def save_processed_data(org_name: str, data: List[dict]):
    """Save processed data for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Convert data to DataFrame
        df = pd.DataFrame(data)
        
        # Save to temp folder
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"pboctotable{org_name_index}{timestamp}"
        
        folder = os.path.join(TEMP_PATH, org_name_index)
        os.makedirs(folder, exist_ok=True)
        
        filepath = os.path.join(folder, f"{filename}.csv")
        df.to_csv(filepath, quoting=1, escapechar='\\')
        
        return {"message": "Data saved successfully", "filename": filename}
    
    except Exception as e:
        logger.error(f"Error saving processed data for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to save processed data")

@router.get("/organizations")
async def get_organizations() -> List[str]:
    """Get list of available organizations"""
    return list(org2name.keys())

class DownloadRequest(BaseModel):
    attachment_ids: List[str]
    force_overwrite: bool = False

class DownloadProgress(BaseModel):
    attachment_id: str
    filename: str
    status: str  # 'pending', 'downloading', 'completed', 'failed', 'skipped'
    progress: int = 0  # 0-100
    downloaded_bytes: int = 0
    total_bytes: int = 0
    speed: str = ""
    error_message: str = ""
    retry_count: int = 0

class DownloadStatus(BaseModel):
    session_id: str
    total_files: int
    completed: int
    failed: int
    skipped: int
    current_file: Optional[str] = None
    overall_progress: int = 0
    files: List[DownloadProgress] = []


class SaveExtractedDataRequest(BaseModel):
    """Payload for saving extracted structured data from the UI."""
    data: List[dict]
    timestamp: Optional[str] = None


@router.post("/save-extracted-data/{org_name}")
async def save_extracted_data(org_name: str, payload: SaveExtractedDataRequest):
    """Save extracted structured penalty data for an organization.

    Accepts a JSON body with `data` (list of dict rows) and optional `timestamp`.
    Persists to a CSV under the temp folder for the given organization, using the
    same naming convention as other saved tables (pboctotable...).
    """
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")

    org_name_index = org2name[org_name]

    try:
        # Validate and normalize data
        rows = payload.data or []

        # Build two datasets: pbocdtl (detail) and pboccat (categorization)
        detail_records: List[Dict[str, Any]] = []
        cat_records: List[Dict[str, Any]] = []

        def norm_date(s: str) -> str:
            if not s:
                return ""
            try:
                # Use pandas to parse various date formats including Chinese formats
                dt = pd.to_datetime(str(s), errors='coerce')
                if pd.isna(dt):
                    return ""
                # If it's a Timestamp with tz, convert to naive date
                return dt.strftime('%Y-%m-%d')
            except Exception:
                return ""

        for r in rows:
            uid = str(uuid.uuid4())
            link = r.get('link', '')
            detail = {
                "企业名称": r.get("被处罚当事人", ""),
                "处罚决定书文号": r.get("行政处罚决定书文号", ""),
                "违法行为类型": r.get("主要违法违规事实", ""),
                "行政处罚依据": r.get("行政处罚依据", ""),
                "行政处罚内容": r.get("行政处罚决定", ""),
                "作出行政处罚决定机关名称": r.get("作出处罚决定的机关名称", ""),
                "作出行政处罚决定日期": r.get("作出处罚决定的日期", ""),
                "link": link,
                "uid": uid,
                "date": norm_date(r.get("作出处罚决定的日期", "")),
            }
            detail_records.append(detail)

            cat = {
                "amount": r.get("罚没总金额", ""),
                "category": r.get("违规类型", ""),
                "province": r.get("监管地区", ""),
                "industry": r.get("行业", ""),
                "id": link,
                "uid": uid,
            }
            cat_records.append(cat)

        df_detail = pd.DataFrame(detail_records)
        df_cat = pd.DataFrame(cat_records)

        # Determine filename
        ts = payload.timestamp
        if ts:
            # Sanitize incoming timestamp: keep digits only and ensure 14 chars (YYYYMMDDHHMMSS)
            ts_digits = ''.join(ch for ch in ts if ch.isdigit())
            if len(ts_digits) >= 14:
                ts = ts_digits[:14]
            else:
                ts = datetime.now().strftime("%Y%m%d%H%M%S")
        else:
            ts = datetime.now().strftime("%Y%m%d%H%M%S")
        # Save as pboc detail dataset under ../pboc
        filename_dtl = f"pbocdtl{org_name_index}{ts}"
        filename_cat = f"pboccat{org_name_index}{ts}"

        # Ensure folder exists and save CSV in PBOC data path (no subfolder)
        os.makedirs(PBOC_DATA_PATH, exist_ok=True)
        filepath_dtl = os.path.join(PBOC_DATA_PATH, f"{filename_dtl}.csv")
        filepath_cat = os.path.join(PBOC_DATA_PATH, f"{filename_cat}.csv")

        # Align with savedf behavior (no special quoting) for pboc dataset
        df_detail.to_csv(filepath_dtl)
        df_cat.to_csv(filepath_cat)

        logger.info(f"Saved extracted data to {filepath_dtl} and {filepath_cat}")
        return {
            "message": "Extracted data saved successfully",
            "filename": filename_dtl,
            "detail_filename": filename_dtl,
            "detail_filepath": filepath_dtl,
            "cat_filename": filename_cat,
            "cat_filepath": filepath_cat,
            "total_records": len(rows),
        }
    except Exception as e:
        logger.error(f"Error saving extracted data for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to save extracted data")

# Global download status storage
download_sessions: Dict[str, DownloadStatus] = {}

async def perform_downloads_background(
    session_id: str,
    selected_attachments: list,
    downloads_dir: str,
    force_overwrite: bool
):
    """Background task to perform the actual downloads"""
    download_status = download_sessions.get(session_id)
    if not download_status:
        logger.error(f"[BACKGROUND_DOWNLOAD] session_id={session_id} Session not found")
        return
    
    downloaded_files = []
    failed_downloads = []
    skipped_files = []
    
    logger.info(f"[BACKGROUND_DOWNLOAD] session_id={session_id} Starting download loop for {len(selected_attachments)} attachments")
    
    for i, attachment in enumerate(selected_attachments):
        file_progress = download_status.files[i]
        download_status.current_file = file_progress.filename
        
        try:
            if not attachment['download_url']:
                file_progress.status = 'failed'
                file_progress.error_message = 'No download URL available'
                failed_downloads.append({
                    'id': attachment['id'],
                    'error': 'No download URL available'
                })
                download_status.failed += 1
                continue
            
            # Generate safe filename
            filename = attachment['filename'] or f"attachment_{attachment['id']}"
            import re
            filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            
            filepath = os.path.join(downloads_dir, filename)
            
            # Check if file already exists
            if os.path.exists(filepath) and not force_overwrite:
                file_progress.status = 'skipped'
                file_progress.progress = 100
                skipped_files.append({
                    'id': attachment['id'],
                    'filename': filename,
                    'reason': 'File already exists'
                })
                download_status.skipped += 1
                continue
            
            # Handle duplicate filenames only if force_overwrite is False
            if not force_overwrite:
                counter = 1
                original_filepath = filepath
                while os.path.exists(filepath):
                    name, ext = os.path.splitext(original_filepath)
                    filepath = f"{name}_{counter}{ext}"
                    counter += 1
                    filename = os.path.basename(filepath)
                    file_progress.filename = filename
            
            # Update progress callback
            def progress_callback(progress_data):
                file_progress.status = 'downloading'
                file_progress.progress = progress_data['progress']
                file_progress.downloaded_bytes = progress_data['downloaded_bytes']
                file_progress.total_bytes = progress_data['total_bytes']
                file_progress.speed = progress_data['speed']
                file_progress.retry_count = progress_data['retry_count']
                
                # Update overall progress
                completed_files = download_status.completed + download_status.failed + download_status.skipped
                current_file_progress = progress_data['progress'] / 100
                download_status.overall_progress = int(((completed_files + current_file_progress) / download_status.total_files) * 100)
                
                # Debug log
                logger.debug(f"[PROGRESS_CALLBACK] session_id={session_id} attachment_id={attachment['id']} progress={progress_data['progress']}% overall={download_status.overall_progress}%")
            
            # Download the file
            logger.info(f"[BACKGROUND_DOWNLOAD] session_id={session_id} attachment_id={attachment['id']} Starting download: {attachment['download_url']}")
            result = download_file_with_progress(
                attachment['download_url'],
                filepath,
                progress_callback,
                max_retries=3,
                referer=attachment.get('referer') or attachment.get('download_url'),
                attachment_id=attachment['id']
            )
            
            if result['success']:
                file_progress.status = 'completed'
                file_progress.progress = 100
                file_progress.retry_count = result['retry_count']
                
                downloaded_files.append({
                    'id': attachment['id'],
                    'filename': os.path.basename(filepath),
                    'filepath': filepath,
                    'size': result['size']
                })
                download_status.completed += 1
                logger.info(f"[BACKGROUND_DOWNLOAD] session_id={session_id} attachment_id={attachment['id']} SUCCESS size={format_bytes(result['size'])} retries={result['retry_count']}")
            else:
                file_progress.status = 'failed'
                file_progress.error_message = result['error']
                file_progress.retry_count = result['retry_count']
                
                failed_downloads.append({
                    'id': attachment['id'],
                    'error': result['error']
                })
                download_status.failed += 1
                logger.error(f"[BACKGROUND_DOWNLOAD] session_id={session_id} attachment_id={attachment['id']} FAILED error={result['error']} retries={result['retry_count']}")
            
        except Exception as e:
            file_progress.status = 'failed'
            file_progress.error_message = str(e)
            
            logger.error(f"[BACKGROUND_DOWNLOAD] session_id={session_id} attachment_id={attachment['id']} EXCEPTION error={str(e)}")
            failed_downloads.append({
                'id': attachment['id'],
                'error': str(e)
            })
            download_status.failed += 1
    
    # Final progress update
    download_status.overall_progress = 100
    download_status.current_file = None
    
    # Log final session summary
    logger.info(f"[BACKGROUND_DOWNLOAD] session_id={session_id} COMPLETED total={len(selected_attachments)} success={len(downloaded_files)} failed={len(failed_downloads)} skipped={len(skipped_files)}")

def download_file_with_progress(
    url: str,
    filepath: str,
    progress_callback=None,
    max_retries: int = 3,
    referer: Optional[str] = None,
    attachment_id: Optional[str] = None,
):
    """Download a file with progress tracking and robust retry logic.

    Changes from prior version:
    - Avoid a preliminary HEAD request (many PBOC sites do not support HEAD).
    - Send a realistic browser header set including optional Referer.
    - Increase read timeouts to accommodate slow servers.
    - Fallback to HTTPS if HTTP fails due to timeouts/disconnects.
    - Enhanced logging with attachment ID for better tracking.
    """
    session = create_session_with_retry()
    filename = os.path.basename(filepath)

    def build_headers(target_url: str) -> Dict[str, str]:
        parsed = urlparse(target_url)
        hdrs = {
            'User-Agent': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                'AppleWebKit/537.36 (KHTML, like Gecko) '
                'Chrome/124.0.0.0 Safari/537.36'
            ),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            # Identity to avoid gzip/br issues on some legacy servers
            'Accept-Encoding': 'identity',
            'Connection': 'keep-alive',
            'Cache-Control': 'no-cache',
            'Pragma': 'no-cache',
            'Host': parsed.netloc,
        }
        if referer:
            hdrs['Referer'] = referer
        return hdrs

    def try_download(target_url: str, attempt_index: int) -> Dict[str, Any]:
        headers = build_headers(target_url)
        # Use separate connect/read timeouts: slower read for large files
        timeout = (15, 120)
        
        # Enhanced logging with attachment ID
        log_prefix = f"[DOWNLOAD] attachment_id={attachment_id or 'unknown'} file={filename}"
        logger.info(f"{log_prefix} Starting download attempt {attempt_index + 1}/{max_retries} url={target_url}")
        
        with session.get(target_url, headers=headers, stream=True, allow_redirects=True, timeout=timeout) as response:
            response.raise_for_status()
            total_size = int(response.headers.get('Content-Length', 0))
            downloaded_size = 0
            start_time = time.time()
            chunk_size = 65536  # 64KB
            last_log_time = start_time
            
            logger.info(f"{log_prefix} Response received, total_size={format_bytes(total_size) if total_size > 0 else 'unknown'}")
            
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if not chunk:
                        continue
                    f.write(chunk)
                    downloaded_size += len(chunk)
                    
                    # Progress callback
                    if progress_callback:
                        elapsed_time = max(time.time() - start_time, 1e-6)
                        speed = downloaded_size / elapsed_time
                        progress = int((downloaded_size / total_size) * 100) if total_size > 0 else 0
                        progress_callback({
                            'downloaded_bytes': downloaded_size,
                            'total_bytes': total_size,
                            'progress': progress,
                            'speed': format_speed(speed),
                            'retry_count': attempt_index,
                        })
                    
                    # Log progress every 5 seconds for large files
                    current_time = time.time()
                    if current_time - last_log_time >= 5.0:
                        elapsed_time = max(current_time - start_time, 1e-6)
                        speed = downloaded_size / elapsed_time
                        progress = int((downloaded_size / total_size) * 100) if total_size > 0 else 0
                        logger.info(f"{log_prefix} Progress: {progress}% ({format_bytes(downloaded_size)}/{format_bytes(total_size) if total_size > 0 else 'unknown'}) speed={format_speed(speed)}")
                        last_log_time = current_time
                        
        logger.info(f"{log_prefix} Download completed successfully, final_size={format_bytes(downloaded_size)}")
        return {
            'success': True,
            'size': downloaded_size,
            'retry_count': attempt_index,
        }

    # Attempt loop with HTTPS fallback if initial scheme is HTTP
    parsed = urlparse(url)
    https_fallback = (parsed.scheme == 'http')
    last_error: Optional[Exception] = None
    for retry_count in range(max_retries):
        try:
            return try_download(url, retry_count)
        except (requests.exceptions.ReadTimeout, requests.exceptions.ConnectTimeout, requests.exceptions.ConnectionError) as e:
            last_error = e
            logger.warning(f"Download attempt {retry_count + 1} failed for {url}: {e}")
            # One-time HTTPS upgrade fallback if using HTTP
            if https_fallback:
                https_url = url.replace('http://', 'https://', 1)
                try:
                    logger.info(f"Retrying via HTTPS fallback: {https_url}")
                    result = try_download(https_url, retry_count)
                    return result
                except Exception as ee:
                    last_error = ee
                    logger.warning(f"HTTPS fallback failed for {https_url}: {ee}")
                    # Only try fallback once
                    https_fallback = False
            # Backoff before next attempt
            time.sleep(2 ** retry_count)
        except requests.exceptions.RequestException as e:
            last_error = e
            logger.warning(f"Download attempt {retry_count + 1} failed for {url}: {e}")
            time.sleep(2 ** retry_count)
        except Exception as e:
            last_error = e
            logger.error(f"Unexpected error downloading {url}: {e}")
            break

    err_msg = f"Network error after {max_retries} retries: {last_error}" if last_error else "Max retries exceeded"
    logger.error(f"All download attempts failed for {url}: {err_msg}")
    return {
        'success': False,
        'error': err_msg,
        'retry_count': max_retries,
    }

@router.post("/download/{org_name}")
async def download_attachments(org_name: str, request: DownloadRequest, background_tasks: BackgroundTasks):
    """Download selected attachments for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    logger.info(f"[DOWNLOAD_REQUEST] org={org_name} attachment_ids={request.attachment_ids} force_overwrite={request.force_overwrite}")
    
    try:
        # Get the download list
        beginwith = f"pboctodownload{org_name_index}"
        
        # First try the organization-specific subfolder
        org_temp_path = os.path.join(TEMP_PATH, org_name_index)
        df = get_csvdf(org_temp_path, beginwith)
        
        # If not found, try the main temp folder
        if df.empty:
            df = get_csvdf(TEMP_PATH, beginwith)
        
        if df.empty:
            raise HTTPException(status_code=404, detail="No attachments found")
        
        # Filter by selected IDs
        selected_attachments = []
        for attachment_id in request.attachment_ids:
            try:
                idx = int(attachment_id)
                if 0 <= idx < len(df):
                    row = df.iloc[idx]
                    selected_attachments.append({
                        'id': attachment_id,
                        'download_url': row.get('download', ''),
                        'filename': os.path.basename(unquote(row.get('download', ''))),
                        # Use the source page as Referer if available; many sites require it
                        'referer': row.get('link', '')
                    })
            except (ValueError, IndexError):
                continue
        
        if not selected_attachments:
            logger.warning(f"[DOWNLOAD_REQUEST] org={org_name} No valid attachments found for IDs: {request.attachment_ids}")
            raise HTTPException(status_code=400, detail="No valid attachments selected")
        
        # Create downloads directory
        downloads_dir = os.path.join(TEMP_PATH, org_name_index, "downloads")
        os.makedirs(downloads_dir, exist_ok=True)
        
        # Create download session
        session_id = f"{org_name}_{int(time.time())}"
        download_status = DownloadStatus(
            session_id=session_id,
            total_files=len(selected_attachments),
            completed=0,
            failed=0,
            skipped=0,
            overall_progress=0,
            files=[]
        )
        
        logger.info(f"[DOWNLOAD_SESSION] session_id={session_id} STARTED org={org_name} total_files={len(selected_attachments)} force_overwrite={request.force_overwrite}")
        
        # Initialize progress for each file
        for attachment in selected_attachments:
            filename = attachment['filename'] or f"attachment_{attachment['id']}"
            import re
            filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            
            progress = DownloadProgress(
                attachment_id=attachment['id'],
                filename=filename,
                status='pending',
                progress=0,
                downloaded_bytes=0,
                total_bytes=0
            )
            download_status.files.append(progress)
        
        download_sessions[session_id] = download_status
        
        # Start background download task
        background_tasks.add_task(
            perform_downloads_background,
            session_id,
            selected_attachments,
            downloads_dir,
            request.force_overwrite
        )
        
        # Return immediately with session info
        return {
            "session_id": session_id,
            "message": f"Started downloading {len(selected_attachments)} files",
            "total_requested": len(selected_attachments),
            "download_status": download_status
        }
    
    except Exception as e:
        logger.error(f"Error downloading attachments for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to download attachments")

@router.get("/download-progress/{session_id}")
async def get_download_progress(session_id: str):
    """Get download progress for a session"""
    if session_id not in download_sessions:
        logger.warning(f"[PROGRESS] session_id={session_id} not found in active sessions: {list(download_sessions.keys())}")
        raise HTTPException(status_code=404, detail="Download session not found")
    
    progress = download_sessions[session_id]
    logger.info(f"[PROGRESS] session_id={session_id} overall_progress={progress.overall_progress}% completed={progress.completed}/{progress.total_files}")
    return progress

@router.delete("/download-session/{session_id}")
async def cleanup_download_session(session_id: str):
    """Clean up download session"""
    if session_id in download_sessions:
        del download_sessions[session_id]
        return {"message": "Session cleaned up successfully"}
    raise HTTPException(status_code=404, detail="Download session not found")

@router.get("/download-sessions")
async def list_download_sessions():
    """List all active download sessions"""
    return {
        "sessions": list(download_sessions.keys()),
        "count": len(download_sessions)
    }

@router.get("/download-file/{org_name}/{filename}")
async def download_file(org_name: str, filename: str):
    """Download a specific file"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    downloads_dir = os.path.join(TEMP_PATH, org_name_index, "downloads")
    filepath = os.path.join(downloads_dir, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/octet-stream'
    )

class AttachmentTextItem(BaseModel):
    id: str
    fileName: str
    fileType: str
    link: str
    filePath: str
    status: str = 'pending'
    content: Optional[str] = None
    errorMessage: Optional[str] = None

class TextExtractionRequest(BaseModel):
    attachment_ids: List[str]
    extract_all: bool = False
    use_soffice: bool = False
    use_llm_ocr: bool = False


def encode_image(image_path: str) -> str:
    """Encode image to base64 string"""
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except Exception as e:
        logger.error(f"Failed to encode image {image_path}: {e}")
        return ""


def llm_ocr_text(image_file: str) -> str:
    """Extract text from image using OpenAI Vision API"""
    try:
        if not os.path.exists(image_file):
            logger.error(f"Image file not found: {image_file}")
            return ""
        
        # Check if OpenAI API key is configured
        if not settings.OPENAI_API_KEY:
            logger.warning("OpenAI API key not configured, skipping LLM OCR")
            return ""
        
        # Get the base64 string
        base64_image = encode_image(image_file)
        if not base64_image:
            logger.error(f"Failed to encode image: {image_file}")
            return ""
        
        logger.info(f"Attempting LLM OCR for: {os.path.basename(image_file)}")
        
        # Use OpenAI client
        try:
            from openai import OpenAI
            client = OpenAI(
                api_key=settings.OPENAI_API_KEY,
                base_url=settings.OPENAI_BASE_URL
            )
            
            response = client.chat.completions.create(
                model=settings.OPENAI_VISION_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "请提取图片中的所有文字内容，包括中文和英文。只返回提取的文字，不要添加任何解释。"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0
            )
            
            if response and response.choices and len(response.choices) > 0:
                content = response.choices[0].message.content
                if content and content.strip():
                    logger.info(f"✓ LLM OCR successful, extracted {len(content)} characters")
                    return content.strip()
            
            logger.warning("❌ Empty response from vision model")
            return ""
            
        except Exception as api_error:
            error_msg = str(api_error)
            logger.error(f"❌ Vision API error: {error_msg}")
            
            # Log specific error details for debugging
            if "400" in error_msg:
                logger.info("💡 HTTP 400 - Check model name and message format")
            elif "401" in error_msg:
                logger.info("💡 HTTP 401 - Check API key")
            elif "404" in error_msg:
                logger.info("💡 HTTP 404 - Model not found")
            
            return ""
        
    except Exception as e:
        logger.error(f"❌ LLM OCR processing error: {str(e)}")
        return ""


def pdfurl2ocr(pdf_path: str, upload_path: str) -> str:
    """Convert PDF to images and extract text using LLM OCR"""
    image_file_list = []
    text = ""
    
    try:
        # Try PyMuPDF first (no external dependencies)
        try:
            import fitz
            logger.info(f"Using PyMuPDF for PDF to image conversion")
            
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            logger.info(f"PDF has {page_count} pages")
            
            # Convert each page to image
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                # Get pixmap with higher resolution for better OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                filename = os.path.join(upload_path, f"page_{page_num + 1}.png")
                pix.save(filename)
                image_file_list.append(filename)
                logger.info(f"Converted page {page_num + 1} to {filename}")
            
            doc.close()
            
        except ImportError:
            logger.info("PyMuPDF not available, trying pdf2image...")
            # Fallback to pdf2image (requires poppler)
            try:
                from pdf2image import convert_from_path
                PDF_file = Path(pdf_path)
                pdf_pages = convert_from_path(PDF_file, 300)  # 300 DPI for good quality
                
                for page_enumeration, page in enumerate(pdf_pages, start=1):
                    filename = os.path.join(upload_path, f"page_{page_enumeration}.jpg")
                    page.save(filename, "JPEG")
                    image_file_list.append(filename)
                    logger.info(f"Converted page {page_enumeration} to {filename}")
                    
            except Exception as pdf2image_error:
                logger.error(f"pdf2image failed: {str(pdf2image_error)}")
                return ""

        # Extract text from images using LLM OCR
        logger.info(f"Starting LLM OCR for {len(image_file_list)} images...")
        for i, image_file in enumerate(image_file_list, 1):
            try:
                logger.info(f"Processing image {i}/{len(image_file_list)}: {os.path.basename(image_file)}")
                extracted_text = llm_ocr_text(image_file)
                if extracted_text and extracted_text.strip():
                    text += extracted_text + "\n"
                    logger.info(f"✓ Extracted {len(extracted_text)} characters from page {i}")
                else:
                    logger.warning(f"❌ No text extracted from page {i}")
            except Exception as e:
                logger.error(f"Error extracting text from {image_file}: {str(e)}")
            finally:
                # Clean up image file
                try:
                    os.remove(image_file)
                except OSError:
                    pass

        if text.strip():
            logger.info(f"✅ Total LLM OCR extraction: {len(text)} characters")
        else:
            logger.warning("❌ No text extracted from any pages")

    except Exception as e:
        logger.error(f"Error in PDF LLM OCR processing for {pdf_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        text = ""
    
    return text.strip()


def extract_text_from_file(file_path: str, file_type: str, use_soffice: bool = False, use_llm_ocr: bool = False) -> str:
    """Extract text content from various file types"""
    try:
        # If use_llm_ocr is enabled, force all files to PDF->Image->LLM OCR pipeline
        if use_llm_ocr:
            logger.info(f"Using LLM OCR mode for {file_type} file: {os.path.basename(file_path)}")
            logger.info(f"File type check: file_type='{file_type}', file_path.lower().endswith('.docx')={file_path.lower().endswith('.docx')}")
            try:
                # Step 1: Convert any file to PDF first (if not already PDF)
                pdf_file_path = file_path
                temp_dir = None
                
                if file_type != 'pdf':
                    # Create temporary directory for conversion
                    temp_dir = tempfile.mkdtemp()
                    
                    try:
                        # Convert file to PDF using soffice
                        logger.info(f"Converting {file_path} to PDF in {temp_dir}")
                        result = subprocess.run([
                            "soffice",
                            "--headless",
                            "--convert-to",
                            "pdf",
                            file_path,
                            "--outdir",
                            temp_dir
                        ], capture_output=True, text=True)
                        
                        logger.info(f"soffice stdout: {result.stdout}")
                        logger.info(f"soffice stderr: {result.stderr}")
                        logger.info(f"soffice return code: {result.returncode}")
                        
                        # Check if soffice conversion was successful
                        if result.returncode != 0:
                            logger.error(f"soffice conversion failed with return code {result.returncode}")
                            logger.error(f"soffice stderr: {result.stderr}")
                            # For docx files, try direct text extraction as fallback
                            if file_type == 'word' and file_path.lower().endswith('.docx'):
                                logger.info("Trying direct docx text extraction as fallback")
                                try:
                                    from docx import Document
                                    doc = Document(file_path)
                                    text = ""
                                    for paragraph in doc.paragraphs:
                                        text += paragraph.text + "\n"
                                    return text.strip() if text.strip() else "DOCX文件内容为空"
                                except Exception as docx_error:
                                    logger.error(f"Direct docx extraction also failed: {docx_error}")
                                    return f"DOCX文件处理失败: soffice转换失败且直接读取也失败 {os.path.basename(file_path)}"
                            return f"{file_type.upper()}文件转PDF失败: soffice无法加载文件 {os.path.basename(file_path)}"
                        
                        # Find the converted PDF file
                        base_name = os.path.splitext(os.path.basename(file_path))[0]
                        pdf_file_path = os.path.join(temp_dir, f"{base_name}.pdf")
                        
                        # List all files in temp directory for debugging
                        temp_files = os.listdir(temp_dir)
                        logger.info(f"Files in temp directory: {temp_files}")
                        
                        if not os.path.exists(pdf_file_path):
                            logger.error(f"Conversion failed: converted file not found at {pdf_file_path}")
                            # For docx files, try direct text extraction as fallback
                            if file_type == 'word' and file_path.lower().endswith('.docx'):
                                logger.info("PDF not found, trying direct docx text extraction as fallback")
                                try:
                                    from docx import Document
                                    # Convert relative path to absolute path
                                    abs_file_path = os.path.abspath(file_path)
                                    doc = Document(abs_file_path)
                                    text = ""
                                    for paragraph in doc.paragraphs:
                                        text += paragraph.text + "\n"
                                    return text.strip() if text.strip() else "DOCX文件内容为空"
                                except Exception as docx_error:
                                    logger.error(f"Direct docx extraction also failed: {docx_error}")
                                    if "Package not found" in str(docx_error) or "PackageNotFoundError" in str(docx_error):
                                        return f"DOCX文件已损坏或格式不正确，无法读取: {os.path.basename(file_path)}"
                                    else:
                                        return f"DOCX文件处理失败: PDF转换失败且直接读取也失败 {os.path.basename(file_path)} (错误: {str(docx_error)})"
                            return f"{file_type.upper()}文件转PDF失败: 转换后的PDF文件未找到"
                    except Exception as e:
                        logger.error(f"LLM OCR mode conversion failed: {e}")
                        return f"{file_type.upper()}文件LLM OCR处理失败: {os.path.basename(file_path)} (转换错误: {str(e)})"
                
                # Step 2: Convert PDF to images and extract text using LLM OCR
                # Create a temporary directory for image conversion if not already created
                if temp_dir is None:
                    temp_dir = tempfile.mkdtemp()
                    
                text = pdfurl2ocr(pdf_file_path, temp_dir)
                
                # Clean up temp directory if created
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
                
                return text
                
            except Exception as e:
                logger.error(f"LLM OCR mode error: {e}")
                return f"{file_type.upper()}文件LLM OCR处理失败: {os.path.basename(file_path)} (错误: {str(e)})"
        
        # If use_soffice is enabled, convert any file to PDF first using soffice
        elif use_soffice:
            logger.info(f"Converting {file_type} file to PDF using soffice: {os.path.basename(file_path)}")
            try:
                
                # Create temporary directory for conversion
                temp_dir = tempfile.mkdtemp()
                
                try:
                    # Convert file to PDF using soffice
                    subprocess.run([
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        file_path,
                        "--outdir",
                        temp_dir
                    ], check=True, capture_output=True, text=True)
                    
                    # Find the converted PDF file
                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                    converted_file = os.path.join(temp_dir, f"{base_name}.pdf")
                    
                    if os.path.exists(converted_file):
                        # Extract text from the converted PDF
                        text = extract_text_from_file(converted_file, 'pdf', False, False)
                        return text
                    else:
                        logger.error(f"Conversion failed: converted file not found at {converted_file}")
                        return f"{file_type.upper()}文件转换失败: {os.path.basename(file_path)}"
                
                except subprocess.CalledProcessError as e:
                    logger.error(f"soffice conversion failed: {e}")
                    return f"{file_type.upper()}文件转换失败: {os.path.basename(file_path)} (soffice转换错误: {e.stderr})"
                except Exception as e:
                    logger.error(f"{file_type.upper()} conversion error: {e}")
                    return f"{file_type.upper()}文件转换失败: {os.path.basename(file_path)} (错误: {str(e)})"
                finally:
                    # Clean up temp directory if it still exists
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir, ignore_errors=True)
                        
            except ImportError:
                logger.warning("subprocess not available, using fallback")
                return f"{file_type.upper()}文件: {os.path.basename(file_path)} (需要安装相关依赖进行转换)"
        
        if file_type == 'pdf':
            # PDF text extraction
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    
                    # If normal PDF extraction failed or returned empty text, try LLM OCR
                    if not text.strip():
                        logger.info("Normal PDF extraction failed, trying LLM OCR...")
                        # Create temporary directory for image conversion
                        temp_dir = tempfile.mkdtemp()
                        try:
                            text = pdfurl2ocr(file_path, temp_dir)
                        finally:
                             # Clean up temporary directory
                             try:
                                 shutil.rmtree(temp_dir)
                             except OSError:
                                 pass
                    
                    return text.strip()
            except ImportError:
                logger.warning("PyPDF2 not available, using fallback for PDF")
                return f"PDF文件: {os.path.basename(file_path)} (需要安装PyPDF2库进行文本提取)"
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {e}")
                # Try LLM OCR as fallback for any PDF processing errors
                logger.info("PDF processing failed, trying LLM OCR as fallback...")
                temp_dir = tempfile.mkdtemp()
                try:
                    text = pdfurl2ocr(file_path, temp_dir)
                    return text.strip()
                except Exception as ocr_error:
                    logger.error(f"LLM OCR fallback also failed: {ocr_error}")
                    return ""
                finally:
                     # Clean up temporary directory
                     try:
                         shutil.rmtree(temp_dir)
                     except OSError:
                         pass
        
        elif file_type == 'word':
            # Word document text extraction - convert all Word formats to PDF first
            try:
                # Convert all Word document formats (.doc, .wps, .docx, .docm) to PDF using soffice
                file_ext = os.path.splitext(file_path)[1].lower()
                logger.info(f"Converting {file_ext} file to PDF: {file_path}")
                
                # Create a temporary directory for conversion
                temp_dir = tempfile.mkdtemp()
                try:
                    # Run soffice conversion to PDF
                    subprocess.run([
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        file_path,
                        "--outdir",
                        temp_dir
                    ], check=True, capture_output=True, text=True)
                    
                    # Find the converted PDF file
                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                    converted_file = os.path.join(temp_dir, f"{base_name}.pdf")
                    
                    if os.path.exists(converted_file):
                        # Extract text from the converted PDF file
                        try:
                            import PyPDF2
                            with open(converted_file, 'rb') as pdf_file:
                                reader = PyPDF2.PdfReader(pdf_file)
                                text = ""
                                for page in reader.pages:
                                    text += page.extract_text() + "\n"
                            
                            # Clean up temporary file
                            os.remove(converted_file)
                            os.rmdir(temp_dir)
                            
                            return text.strip()
                        except ImportError:
                            logger.warning("PyPDF2 not available for PDF text extraction")
                            return f"{file_ext.upper()}文件已转换为PDF，但需要安装PyPDF2库进行文本提取: {os.path.basename(file_path)}"
                    else:
                        logger.error(f"Conversion failed: converted file not found at {converted_file}")
                        return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)}"
                
                except subprocess.CalledProcessError as e:
                    logger.error(f"soffice conversion failed: {e}")
                    return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)} (soffice转换错误: {e.stderr})"
                except Exception as e:
                    logger.error(f"{file_ext.upper()} conversion error: {e}")
                    return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)} (错误: {str(e)})"
                finally:
                    # Clean up temp directory if it still exists
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir, ignore_errors=True)
                    
            except ImportError:
                logger.warning("subprocess not available, using fallback for Word")
                return f"Word文档: {os.path.basename(file_path)} (需要安装相关依赖进行文本提取)"
        
        elif file_type == 'excel':
            # Excel text extraction - convert to PDF first
            try:
                # Convert Excel files (.xls, .xlsx) to PDF using soffice
                file_ext = os.path.splitext(file_path)[1].lower()
                logger.info(f"Converting {file_ext} file to PDF: {file_path}")
                
                # Create a temporary directory for conversion
                temp_dir = tempfile.mkdtemp()
                try:
                    # Run soffice conversion to PDF
                    subprocess.run([
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        file_path,
                        "--outdir",
                        temp_dir
                    ], check=True, capture_output=True, text=True)
                    
                    # Find the converted PDF file
                    base_name = os.path.splitext(os.path.basename(file_path))[0]
                    converted_file = os.path.join(temp_dir, f"{base_name}.pdf")
                    
                    if os.path.exists(converted_file):
                        logger.info(f"Successfully converted to PDF: {converted_file}")
                        
                        # Extract text from the converted PDF
                        try:
                            import PyPDF2
                            with open(converted_file, 'rb') as pdf_file:
                                reader = PyPDF2.PdfReader(pdf_file)
                                text = ""
                                for page in reader.pages:
                                    text += page.extract_text() + "\n"
                                return text.strip()
                        except ImportError:
                            logger.warning("PyPDF2 not available for Excel PDF conversion")
                            return f"{file_ext.upper()}文件已转换为PDF，但需要安装PyPDF2库进行文本提取: {os.path.basename(file_path)}"
                    else:
                        logger.error(f"Conversion failed: converted file not found at {converted_file}")
                        return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)}"
                
                except subprocess.CalledProcessError as e:
                    logger.error(f"soffice conversion failed: {e}")
                    return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)} (soffice转换错误: {e.stderr})"
                except Exception as e:
                    logger.error(f"{file_ext.upper()} conversion error: {e}")
                    return f"{file_ext.upper()}文件转换失败: {os.path.basename(file_path)} (错误: {str(e)})"
                finally:
                    # Clean up temp directory if it still exists
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir, ignore_errors=True)
                    
            except ImportError:
                logger.warning("subprocess not available, using fallback for Excel")
                return f"Excel文件: {os.path.basename(file_path)} (需要安装相关依赖进行文本提取)"
        
        elif file_type == 'image':
            # Image OCR text extraction
            try:
                import pytesseract
                from PIL import Image
                image = Image.open(file_path)
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                return text.strip()
            except ImportError:
                logger.warning("pytesseract or PIL not available, using fallback for image")
                return f"图片文件: {os.path.basename(file_path)} (需要安装pytesseract和PIL库进行OCR文本提取)"
        
        else:
            # Try to read as plain text
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except UnicodeDecodeError:
                try:
                    with open(file_path, 'r', encoding='gbk') as file:
                        return file.read()
                except UnicodeDecodeError:
                    return f"无法读取文件: {os.path.basename(file_path)} (编码格式不支持)"
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return f"文本提取失败: {str(e)}"

@router.get("/attachment-text-list/{org_name}")
async def get_attachment_text_list(org_name: str) -> List[AttachmentTextItem]:
    """Get list of attachments with text extraction capability for an organization"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Look for download files in temp folder
        beginwith = f"pboctodownload{org_name_index}"
        
        # First try the organization-specific subfolder
        org_temp_path = os.path.join(TEMP_PATH, org_name_index)
        df = get_csvdf(org_temp_path, beginwith)
        
        # If not found, try the main temp folder
        if df.empty:
            df = get_csvdf(TEMP_PATH, beginwith)
        
        if df.empty:
            return []
        
        # Check downloads directory for existing files
        downloads_dir = os.path.join(TEMP_PATH, org_name_index, "downloads")
        if not os.path.exists(downloads_dir):
            return []
        
        existing_files = set(os.listdir(downloads_dir))
        
        attachments = []
        for idx, row in df.iterrows():
            download_url = row.get('download', '')
            file_name = os.path.basename(download_url) if download_url else f"file_{idx}"
            # Decode URL-encoded filename
            file_name = unquote(file_name)
            
            # Clean filename for comparison
            import re
            clean_filename = re.sub(r'[<>:"/\\|?*]', '_', file_name)
            
            # Check if file exists in downloads directory
            file_path = None
            for existing_file in existing_files:
                if existing_file == clean_filename or existing_file.startswith(clean_filename.split('.')[0]):
                    file_path = os.path.join(downloads_dir, existing_file)
                    clean_filename = existing_file
                    break
            
            if file_path and os.path.exists(file_path):
                # Determine file type
                file_ext = os.path.splitext(clean_filename)[1].lower()
                if file_ext in ['.xlsx', '.xls', '.et', '.ett']:
                    file_type = 'excel'
                elif file_ext == '.pdf':
                    file_type = 'pdf'
                elif file_ext in ['.docx', '.doc', '.wps', '.docm']:
                    file_type = 'word'
                elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tif']:
                    file_type = 'image'
                else:
                    file_type = 'other'
                
                attachment = AttachmentTextItem(
                    id=str(idx),
                    fileName=clean_filename,
                    fileType=file_type,
                    link=row.get('link', ''),
                    filePath=file_path,
                    status='ready'
                )
                attachments.append(attachment)
        
        return attachments
    
    except Exception as e:
        logger.error(f"Error getting attachment text list for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get attachment text list")

@router.post("/extract-text/{org_name}")
async def extract_attachment_text(org_name: str, request: TextExtractionRequest):
    """Extract text from selected attachments"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Get attachment list
        attachments = await get_attachment_text_list(org_name)
        
        if not attachments:
            raise HTTPException(status_code=404, detail="No attachments found")
        
        # Filter by selected IDs or extract all
        if request.extract_all:
            selected_attachments = attachments
        else:
            selected_attachments = [att for att in attachments if att.id in request.attachment_ids]
        
        if not selected_attachments:
            raise HTTPException(status_code=400, detail="No valid attachments selected")
        
        # Extract text from each file
        results = []
        for attachment in selected_attachments:
            try:
                logger.info(f"Extracting text from {attachment.fileName}")
                content = extract_text_from_file(attachment.filePath, attachment.fileType, request.use_soffice, request.use_llm_ocr)
                
                result = {
                    "id": attachment.id,
                    "fileName": attachment.fileName,
                    "fileType": attachment.fileType,
                    "link": attachment.link,
                    "filePath": attachment.filePath,
                    "status": "completed",
                    "content": content
                }
                results.append(result)
                
            except Exception as e:
                logger.error(f"Error extracting text from {attachment.fileName}: {e}")
                result = {
                    "id": attachment.id,
                    "fileName": attachment.fileName,
                    "fileType": attachment.fileType,
                    "link": attachment.link,
                    "filePath": attachment.filePath,
                    "status": "failed",
                    "content": None,
                    "errorMessage": str(e)
                }
                results.append(result)
        
        return {
            "message": f"Text extraction completed for {len(results)} files",
            "total_files": len(results),
            "successful": len([r for r in results if r["status"] == "completed"]),
            "failed": len([r for r in results if r["status"] == "failed"]),
            "results": results
        }
    
    except Exception as e:
        logger.error(f"Error extracting text for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract text")

@router.post("/save-text-results/{org_name}")
async def save_text_results(org_name: str, results: List[dict]):
    """Save text extraction results as pboctotable file"""
    if org_name not in org2name:
        raise HTTPException(status_code=400, detail="Invalid organization name")
    
    org_name_index = org2name[org_name]
    
    try:
        # Convert results to DataFrame
        df = pd.DataFrame(results)
        
        # Save to temp folder
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"pboctotable{org_name_index}{timestamp}"
        
        folder = os.path.join(TEMP_PATH, org_name_index)
        os.makedirs(folder, exist_ok=True)
        
        filepath = os.path.join(folder, f"{filename}.csv")
        df.to_csv(filepath, quoting=1, escapechar='\\')
        
        logger.info(f"Saved text extraction results to {filepath}")
        
        return {
            "message": "Text extraction results saved successfully",
            "filename": filename,
            "filepath": filepath,
            "total_records": len(results)
        }
    
    except Exception as e:
        logger.error(f"Error saving text results for {org_name}: {e}")
        raise HTTPException(status_code=500, detail="Failed to save text results")
