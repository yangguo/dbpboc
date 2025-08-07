from typing import List, Optional
from datetime import datetime
import os
import uuid
import aiofiles
from fastapi import UploadFile
from motor.motor_asyncio import AsyncIOMotorDatabase
from app.models.document import Document, DocumentCreate, DocumentUpdate
from app.core.config import settings
from bson import ObjectId
import mimetypes

class DocumentService:
    def __init__(self, database: AsyncIOMotorDatabase):
        self.db = database
        self.collection = database.documents
        
        # Ensure upload directory exists
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    
    async def upload_document(
        self,
        file: UploadFile,
        document_type: str,
        case_id: Optional[str] = None,
        uploaded_by: Optional[str] = None
    ) -> Document:
        """Upload and store a document"""
        # Validate file size
        if file.size and file.size > settings.MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds maximum allowed size of {settings.MAX_FILE_SIZE} bytes")
        
        # Generate unique filename
        file_extension = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
        
        # Save file to disk
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        # Get file info
        file_size = len(content)
        mime_type = mimetypes.guess_type(file.filename)[0] or "application/octet-stream"
        
        # Create document record
        document_data = DocumentCreate(
            filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type,
            case_id=ObjectId(case_id) if case_id else None,
            document_type=document_type,
            processing_status="uploaded"
        )
        
        document_dict = document_data.dict()
        document_dict["created_at"] = datetime.utcnow()
        document_dict["updated_at"] = datetime.utcnow()
        document_dict["uploaded_by"] = uploaded_by
        
        result = await self.collection.insert_one(document_dict)
        document_dict["_id"] = result.inserted_id
        
        document = Document(**document_dict)
        
        # Start background processing if it's a processable document
        if self._is_processable(mime_type):
            # In a real application, you would queue this for background processing
            # For now, we'll process it immediately
            await self._process_document_content(document)
        
        return document
    
    async def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """Get document by ID"""
        document_doc = await self.collection.find_one({"_id": ObjectId(document_id)})
        if document_doc:
            return Document(**document_doc)
        return None
    
    async def list_documents(
        self,
        case_id: Optional[str] = None,
        document_type: Optional[str] = None,
        processing_status: Optional[str] = None,
        page: int = 1,
        size: int = 20
    ) -> List[Document]:
        """List documents with filters"""
        query = {}
        
        if case_id:
            query["case_id"] = ObjectId(case_id)
        
        if document_type:
            query["document_type"] = document_type
        
        if processing_status:
            query["processing_status"] = processing_status
        
        skip = (page - 1) * size
        cursor = self.collection.find(query).skip(skip).limit(size)
        cursor = cursor.sort("created_at", -1)
        
        documents = []
        async for doc in cursor:
            documents.append(Document(**doc))
        
        return documents
    
    async def update_document(self, document_id: str, document_update: DocumentUpdate) -> Optional[Document]:
        """Update document metadata"""
        update_data = document_update.dict(exclude_unset=True)
        update_data["updated_at"] = datetime.utcnow()
        
        result = await self.collection.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": update_data}
        )
        
        if result.modified_count:
            return await self.get_document_by_id(document_id)
        return None
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete document and its file"""
        document = await self.get_document_by_id(document_id)
        if not document:
            return False
        
        # Delete file from disk
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception:
            pass  # Continue even if file deletion fails
        
        # Delete document record
        result = await self.collection.delete_one({"_id": ObjectId(document_id)})
        return result.deleted_count > 0
    
    async def process_document(self, document_id: str) -> bool:
        """Trigger document processing"""
        document = await self.get_document_by_id(document_id)
        if not document:
            raise ValueError("Document not found")
        
        if not self._is_processable(document.mime_type):
            raise ValueError("Document type is not processable")
        
        await self._process_document_content(document)
        return True
    
    def _is_processable(self, mime_type: str) -> bool:
        """Check if document type can be processed"""
        processable_types = [
            "application/pdf",
            "image/jpeg",
            "image/png",
            "image/tiff",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain"
        ]
        return mime_type in processable_types
    
    async def _process_document_content(self, document: Document):
        """Process document content (OCR, text extraction)"""
        try:
            # Update status to processing
            await self.collection.update_one(
                {"_id": document.id},
                {"$set": {"processing_status": "processing", "updated_at": datetime.utcnow()}}
            )
            
            extracted_text = ""
            ocr_confidence = None
            
            if document.mime_type == "application/pdf":
                extracted_text, ocr_confidence = await self._process_pdf(document.file_path)
            elif document.mime_type.startswith("image/"):
                extracted_text, ocr_confidence = await self._process_image(document.file_path)
            elif document.mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                extracted_text = await self._process_docx(document.file_path)
            elif document.mime_type == "text/plain":
                extracted_text = await self._process_text(document.file_path)
            
            # Update document with extracted content
            update_data = {
                "extracted_text": extracted_text,
                "processing_status": "completed",
                "updated_at": datetime.utcnow()
            }
            
            if ocr_confidence is not None:
                update_data["ocr_confidence"] = ocr_confidence
            
            await self.collection.update_one(
                {"_id": document.id},
                {"$set": update_data}
            )
            
        except Exception as e:
            # Update status to failed
            await self.collection.update_one(
                {"_id": document.id},
                {"$set": {
                    "processing_status": "failed",
                    "updated_at": datetime.utcnow(),
                    "metadata.error": str(e)
                }}
            )
    
    async def _process_pdf(self, file_path: str) -> tuple[str, Optional[float]]:
        """Extract text from PDF using pdfplumber"""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip(), None
        except Exception as e:
            # Fallback to OCR if text extraction fails
            return await self._process_image(file_path)
    
    async def _process_image(self, file_path: str) -> tuple[str, Optional[float]]:
        """Extract text from image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            import cv2
            import numpy as np
            
            # Load and preprocess image
            image = cv2.imread(file_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply some preprocessing to improve OCR accuracy
            gray = cv2.medianBlur(gray, 3)
            
            # Convert back to PIL Image for pytesseract
            pil_image = Image.fromarray(gray)
            
            # Extract text with confidence
            data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT)
            
            # Calculate average confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Extract text
            text = pytesseract.image_to_string(pil_image, lang='chi_sim+eng')
            
            return text.strip(), avg_confidence / 100.0  # Convert to 0-1 scale
            
        except Exception as e:
            return f"OCR processing failed: {str(e)}", 0.0
    
    async def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            return f"DOCX processing failed: {str(e)}"
    
    async def _process_text(self, file_path: str) -> str:
        """Read plain text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='gbk') as f:
                    return await f.read()
            except Exception as e:
                return f"Text processing failed: {str(e)}"
        except Exception as e:
            return f"Text processing failed: {str(e)}"
    
    async def get_document_statistics(self) -> dict:
        """Get document statistics"""
        pipeline = [
            {
                "$group": {
                    "_id": None,
                    "total_documents": {"$sum": 1},
                    "total_size": {"$sum": "$file_size"}
                }
            }
        ]
        
        result = await self.collection.aggregate(pipeline).to_list(1)
        stats = result[0] if result else {"total_documents": 0, "total_size": 0}
        
        # Get documents by type
        type_pipeline = [
            {"$group": {"_id": "$document_type", "count": {"$sum": 1}}}
        ]
        type_result = await self.collection.aggregate(type_pipeline).to_list(None)
        stats["by_type"] = {item["_id"]: item["count"] for item in type_result}
        
        # Get processing status
        status_pipeline = [
            {"$group": {"_id": "$processing_status", "count": {"$sum": 1}}}
        ]
        status_result = await self.collection.aggregate(status_pipeline).to_list(None)
        stats["by_status"] = {item["_id"]: item["count"] for item in status_result}
        
        return stats